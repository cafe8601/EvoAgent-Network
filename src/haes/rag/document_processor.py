"""
Document Processor - 문서 청킹 및 전처리

RAG 시스템용 문서 처리:
- 파일 로드 (PDF, TXT, MD, DOCX)
- 텍스트 청킹
- 메타데이터 추출

05-data-processing SKILL 참조
"""

import os
import re
from typing import List, Optional, Dict, Any, Callable
from pathlib import Path
from dataclasses import dataclass
from loguru import logger


@dataclass
class DocumentChunk:
    """문서 청크"""
    content: str
    metadata: Dict[str, Any]
    chunk_index: int
    start_char: int
    end_char: int


@dataclass
class ProcessedDocument:
    """처리된 문서"""
    source_path: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    total_chars: int
    total_chunks: int


class DocumentProcessor:
    """
    문서 프로세서
    
    다양한 파일 포맷을 로드하고 청킹
    
    사용 예시:
        processor = DocumentProcessor(chunk_size=1000, chunk_overlap=200)
        
        # 파일 처리
        doc = processor.process_file("./document.md")
        
        # 텍스트 청킹
        chunks = processor.chunk_text("Long text here...", metadata={"source": "api"})
        
        # 디렉토리 처리
        docs = processor.process_directory("./docs", extensions=[".md", ".txt"])
    """
    
    # 지원 형식
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.markdown', '.py', '.js', '.ts', '.json', '.yaml', '.yml'}
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        length_function: Callable[[str], int] = len,
    ):
        """
        Args:
            chunk_size: 청크 크기 (문자 수)
            chunk_overlap: 청크 간 겹침
            length_function: 길이 계산 함수
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.length_function = length_function
        
        # PDF 라이브러리 체크
        self._pdf_available = self._check_pdf_support()
        self._docx_available = self._check_docx_support()
    
    def _check_pdf_support(self) -> bool:
        """PDF 지원 체크"""
        try:
            import pypdf
            return True
        except ImportError:
            try:
                import PyPDF2
                return True
            except ImportError:
                return False
    
    def _check_docx_support(self) -> bool:
        """DOCX 지원 체크"""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def load_file(self, file_path: str) -> str:
        """
        파일 로드
        
        Args:
            file_path: 파일 경로
        
        Returns:
            파일 텍스트 내용
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            return self._load_pdf(path)
        elif ext == '.docx':
            return self._load_docx(path)
        elif ext in self.SUPPORTED_EXTENSIONS or ext in {'.html', '.htm'}:
            return self._load_text(path)
        else:
            # 기본적으로 텍스트로 시도
            try:
                return self._load_text(path)
            except Exception as e:
                raise ValueError(f"Unsupported file type: {ext}. Error: {e}")
    
    def _load_text(self, path: Path) -> str:
        """텍스트 파일 로드"""
        encodings = ['utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'latin-1']
        
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        
        # 마지막 시도: 바이너리로 읽고 디코드
        return path.read_bytes().decode('utf-8', errors='ignore')
    
    def _load_pdf(self, path: Path) -> str:
        """PDF 파일 로드"""
        if not self._pdf_available:
            raise RuntimeError("PDF support not available. Install: pip install pypdf")
        
        try:
            import pypdf
            
            with open(path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return '\n\n'.join(text_parts)
        except ImportError:
            import PyPDF2
            
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return '\n\n'.join(text_parts)
    
    def _load_docx(self, path: Path) -> str:
        """DOCX 파일 로드"""
        if not self._docx_available:
            raise RuntimeError("DOCX support not available. Install: pip install python-docx")
        
        import docx
        
        doc = docx.Document(str(path))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        return '\n\n'.join(text_parts)
    
    def chunk_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[DocumentChunk]:
        """
        텍스트 청킹
        
        Args:
            text: 원본 텍스트
            metadata: 메타데이터
        
        Returns:
            청크 리스트
        """
        metadata = metadata or {}
        
        if not text.strip():
            return []
        
        chunks = []
        
        # 문단 기반 분할 먼저 시도
        paragraphs = self._split_by_paragraphs(text)
        
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for para in paragraphs:
            # 현재 청크 + 새 문단이 chunk_size 이하면 합침
            if self.length_function(current_chunk + para) <= self.chunk_size:
                current_chunk += para + "\n\n"
            else:
                # 현재 청크 저장
                if current_chunk.strip():
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        metadata=metadata.copy(),
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                    ))
                    chunk_index += 1
                    current_start += len(current_chunk) - self.chunk_overlap
                
                # 새 청크 시작
                if self.length_function(para) > self.chunk_size:
                    # 긴 문단은 문장 단위로 분할
                    sub_chunks = self._split_long_text(para, metadata, chunk_index, current_start)
                    chunks.extend(sub_chunks)
                    chunk_index += len(sub_chunks)
                    current_chunk = ""
                else:
                    # 오버랩 부분 포함
                    if current_chunk:
                        overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                        current_chunk = overlap_text + para + "\n\n"
                    else:
                        current_chunk = para + "\n\n"
        
        # 마지막 청크
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                metadata=metadata.copy(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
            ))
        
        return chunks
    
    def _split_by_paragraphs(self, text: str) -> List[str]:
        """문단 분할"""
        # 빈 줄로 분할
        paragraphs = re.split(r'\n\s*\n', text)
        return [p.strip() for p in paragraphs if p.strip()]
    
    def _split_long_text(
        self,
        text: str,
        metadata: Dict[str, Any],
        start_index: int,
        start_char: int,
    ) -> List[DocumentChunk]:
        """긴 텍스트 분할"""
        chunks = []
        
        # 문장 단위로 분할
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        current_chunk = ""
        chunk_index = start_index
        current_start = start_char
        
        for sentence in sentences:
            if self.length_function(current_chunk + sentence) <= self.chunk_size:
                current_chunk += sentence + " "
            else:
                if current_chunk.strip():
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        metadata=metadata.copy(),
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                    ))
                    chunk_index += 1
                    current_start += len(current_chunk) - self.chunk_overlap
                
                # 오버랩 포함 새 청크
                overlap = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                current_chunk = overlap + sentence + " "
        
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                metadata=metadata.copy(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
            ))
        
        return chunks
    
    def process_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """
        파일 처리 (로드 + 청킹)
        
        Args:
            file_path: 파일 경로
            metadata: 추가 메타데이터
        
        Returns:
            처리된 문서
        """
        path = Path(file_path)
        
        # 파일 로드
        text = self.load_file(file_path)
        
        # 기본 메타데이터
        file_metadata = {
            "source": str(path.absolute()),
            "filename": path.name,
            "extension": path.suffix,
            "size_bytes": path.stat().st_size,
        }
        
        if metadata:
            file_metadata.update(metadata)
        
        # 청킹
        chunks = self.chunk_text(text, file_metadata)
        
        return ProcessedDocument(
            source_path=str(path.absolute()),
            chunks=chunks,
            metadata=file_metadata,
            total_chars=len(text),
            total_chunks=len(chunks),
        )
    
    def process_directory(
        self,
        directory_path: str,
        extensions: Optional[List[str]] = None,
        recursive: bool = True,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[ProcessedDocument]:
        """
        디렉토리 처리
        
        Args:
            directory_path: 디렉토리 경로
            extensions: 처리할 확장자 (None이면 모두)
            recursive: 하위 디렉토리 포함
            metadata: 공통 메타데이터
        
        Returns:
            처리된 문서 리스트
        """
        path = Path(directory_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        if not path.is_dir():
            raise ValueError(f"Not a directory: {directory_path}")
        
        # 확장자 필터
        if extensions:
            extensions = {ext.lower() if ext.startswith('.') else f'.{ext}'.lower() for ext in extensions}
        
        # 파일 수집
        if recursive:
            files = list(path.rglob('*'))
        else:
            files = list(path.glob('*'))
        
        documents = []
        
        for file_path in files:
            if not file_path.is_file():
                continue
            
            if extensions and file_path.suffix.lower() not in extensions:
                continue
            
            try:
                doc = self.process_file(str(file_path), metadata)
                documents.append(doc)
                logger.debug(f"Processed: {file_path} ({doc.total_chunks} chunks)")
            except Exception as e:
                logger.warning(f"Failed to process {file_path}: {e}")
        
        logger.info(f"Processed {len(documents)} files from {directory_path}")
        return documents
    
    def process_skill_directory(
        self,
        skills_path: str,
        skill_ids: Optional[List[str]] = None,
    ) -> List[ProcessedDocument]:
        """
        AI-research-SKILLs 디렉토리 처리
        
        Args:
            skills_path: SKILLs 디렉토리 경로
            skill_ids: 처리할 스킬 ID 목록 (None이면 모두)
        
        Returns:
            처리된 문서 리스트
        """
        path = Path(skills_path)
        documents = []
        
        # 스킬 디렉토리 순회
        for skill_dir in path.iterdir():
            if not skill_dir.is_dir():
                continue
            
            skill_id = skill_dir.name
            
            # 스킬 필터
            if skill_ids and skill_id not in skill_ids:
                continue
            
            # SKILL.md 파일 찾기
            skill_file = skill_dir / "SKILL.md"
            if skill_file.exists():
                try:
                    doc = self.process_file(
                        str(skill_file),
                        metadata={
                            "skill_id": skill_id,
                            "type": "skill",
                        }
                    )
                    documents.append(doc)
                except Exception as e:
                    logger.warning(f"Failed to process {skill_file}: {e}")
            
            # 하위 스킬 처리
            for sub_dir in skill_dir.iterdir():
                if sub_dir.is_dir():
                    sub_skill_file = sub_dir / "SKILL.md"
                    if sub_skill_file.exists():
                        try:
                            doc = self.process_file(
                                str(sub_skill_file),
                                metadata={
                                    "skill_id": f"{skill_id}/{sub_dir.name}",
                                    "parent_skill": skill_id,
                                    "type": "sub_skill",
                                }
                            )
                            documents.append(doc)
                        except Exception as e:
                            logger.warning(f"Failed to process {sub_skill_file}: {e}")
        
        logger.info(f"Processed {len(documents)} skill documents from {skills_path}")
        return documents


# 편의 함수
def chunk_text(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> List[str]:
    """텍스트 청킹 (간단 버전)"""
    processor = DocumentProcessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = processor.chunk_text(text)
    return [c.content for c in chunks]
